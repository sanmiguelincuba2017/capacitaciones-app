import os
os.environ.setdefault('OMP_NUM_THREADS', '1')
os.environ.setdefault('OPENBLAS_NUM_THREADS', '1')
os.environ.setdefault('MKL_NUM_THREADS', '1')

import pandas as pd
import numpy as np
import gc
from datetime import datetime
from flask import Flask, render_template, request, redirect, url_for, flash, send_file, session
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import json
import hashlib

app = Flask(__name__)
app.secret_key = 'capacitaciones_secret_key_2026'

# Configuración
UPLOAD_FOLDER = 'uploads'
GENERATED_FOLDER = 'generated'
DATA_FILE = 'data/capacitaciones.json'
ALLOWED_EXTENSIONS = {'xlsx', 'xls'}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(GENERATED_FOLDER, exist_ok=True)
os.makedirs('data', exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max

# Usuario único (podés cambiar esto)
USUARIO = {
    'username': 'admin',
    'password': hashlib.sha256('admin123'.encode()).hexdigest()
}

# ============================================================
# FUNCIONES DE PROCESAMIENTO (BACKEND VALIDADO)
# ============================================================

CLASIFICACION_CP = {
    'San Miguel': [1659, 1661, 1662, 1663],
    'José C. Paz': [1665, 1666],
    'Bella Vista': [1661],
    'Moreno': [1742, 1743, 1744, 1745, 1746],
    'CABA': list(range(1000, 1500)),
    'Malvinas Argentinas': [1611, 1612, 1613, 1614, 1615, 1616, 1617, 1667],
    'Hurlingham': [1686, 1688],
    'San Martín': [1650, 1651, 1652, 1653, 1655, 1672],
    'Vicente López': [1602, 1603, 1604, 1605, 1606, 1636, 1638],
    'Tigre': [1601, 1618, 1621, 1624, 1648, 1670],
    'General Rodríguez': [1748],
    'Pablo Nogués': [1616],
    'La Matanza': list(range(1750, 1771)),
}

def clasificar_partido(cp):
    if pd.isna(cp):
        return 'Sin CP'
    cp = int(cp)
    if cp == 16663:
        cp = 1663
    for partido, cps in CLASIFICACION_CP.items():
        if cp in cps:
            return partido
    if 1000 <= cp <= 1999:
        return 'Otros del GBA'
    return 'Otras Provincias'

def limpiar_celular(cel):
    if pd.isna(cel):
        return None
    cel = str(cel).replace(' ', '').replace('.', '').replace('-', '').replace('(', '').replace(')', '')
    if cel.lower() in ['nan', 'none', 'no tengo', 'no', '.', '']:
        return None
    if cel.startswith('54'):
        cel = cel[2:]
    if cel.startswith('9'):
        cel = cel[1:]
    if cel.startswith('0'):
        cel = cel[1:]
    cel = ''.join(filter(str.isdigit, cel))
    return cel if cel else None

def limpiar_dni(dni):
    if pd.isna(dni):
        return None
    dni = str(dni).replace('.', '').replace('-', '').replace(' ', '').replace('/', '').replace('_', '')
    if len(dni) > 8:
        dni = dni[-8:]
    return dni if dni.isdigit() else None

def obtener_periodo(fecha_dt, tipo):
    """Devuelve una etiqueta de período (trimestre/semestre/año) para una fecha dada."""
    anio = fecha_dt.year
    mes = fecha_dt.month
    if tipo == 'trimestral':
        trimestre = (mes - 1) // 3 + 1
        return f"{anio}-T{trimestre}"
    if tipo == 'semestral':
        semestre = 1 if mes <= 6 else 2
        return f"{anio}-S{semestre}"
    return str(anio)  # anual


def procesar_capacitacion(filepath_inscriptos, filepath_mat_cert, nombre_cap, inscritos, matriculados_num, certificados_num):
    """Procesa una capacitación completa y genera Excel + Word."""

    # Leer inscriptos
    df_inscriptos_raw = pd.read_excel(filepath_inscriptos, header=None, engine='openpyxl', engine_kwargs={'read_only': True})
    df_ins = df_inscriptos_raw.copy()
    df_ins.columns = ['fecha_inscripcion', 'email', 'capacitaciones', 'nombre', 'apellido', 
                      'rango_edad', 'fecha_nac', 'dni', 'cuil', 'genero', 'calle', 'numero',
                      'col12', 'col13', 'col14', 'localidad_barrio', 'ciudad', 'cp', 'provincia',
                      'celular', 'celular_alt', 'nivel_educativo', 'sector_educativo', 'situacion_laboral',
                      'dispositivos', 'conectividad', 'tipo_negocio', 'etapa_emprendimiento',
                      'tipo_actividad', 'rubro', 'descripcion_emprendimiento', 'tipo_empresa',
                      'nombre_emprendimiento', 'instagram', 'expectativas']

    df_ins['cp'] = pd.to_numeric(df_ins['cp'], errors='coerce')
    df_ins['dni'] = df_ins['dni'].apply(limpiar_dni)
    df_ins['celular'] = df_ins['celular'].apply(limpiar_celular)
    df_ins['celular_alt'] = df_ins['celular_alt'].apply(limpiar_celular)
    df_ins['partido'] = df_ins['cp'].apply(clasificar_partido)
    df_ins['nombre'] = df_ins['nombre'].str.strip()
    df_ins['apellido'] = df_ins['apellido'].str.strip()

    # Leer matriculados/certificados
    df_mat_cert_raw = pd.read_excel(filepath_mat_cert, sheet_name='MATRICULADOCERTIFICADO', header=None, engine='openpyxl', engine_kwargs={'read_only': True})

    # Extraer matriculados
    df_mat = df_mat_cert_raw.iloc[7:, [0, 1, 2, 3, 4, 5]].copy()
    df_mat.columns = ['n_inscripto', 'nombre', 'apellido', 'apellido_nombre', 'dni', 'celular']
    df_mat = df_mat.dropna(subset=['dni'])
    df_mat['dni'] = df_mat['dni'].apply(limpiar_dni)
    df_mat['celular'] = df_mat['celular'].apply(limpiar_celular)

    # Extraer certificados
    df_cert = df_mat_cert_raw.iloc[7:, [12, 13, 14, 15, 16, 17]].copy()
    df_cert.columns = ['n_inscripto', 'nombre', 'apellido', 'celular', 'apellido_nombre', 'presente']
    df_cert = df_cert.dropna(subset=['celular'])
    df_cert['celular'] = df_cert['celular'].apply(limpiar_celular)

    # Cruzar matriculados con certificados
    cert_celulares = set(df_cert['celular'].dropna())
    cert_dnis = set()
    for cel in cert_celulares:
        match = df_ins[df_ins['celular'] == cel]
        if not match.empty and match.iloc[0]['dni']:
            cert_dnis.add(match.iloc[0]['dni'])

    df_mat['certificado'] = df_mat.apply(
        lambda row: 'Sí' if (row['celular'] in cert_celulares or row['dni'] in cert_dnis) else 'No',
        axis=1
    )

    # Generar Excel
    fecha_hoy = datetime.now().strftime('%Y%m%d')
    safe_name = nombre_cap.replace(' ', '_').replace(':', '').replace('/', '')[:50]
    excel_name = f"{fecha_hoy}_{safe_name}.xlsx"
    excel_path = os.path.join(GENERATED_FOLDER, excel_name)

    hoja1 = df_ins[['fecha_inscripcion', 'email', 'nombre', 'apellido', 'dni', 
                    'celular', 'celular_alt', 'calle', 'numero', 'localidad_barrio',
                    'ciudad', 'cp', 'partido', 'provincia', 'rango_edad', 'genero',
                    'nivel_educativo', 'situacion_laboral', 'tipo_actividad', 
                    'rubro', 'nombre_emprendimiento', 'instagram']].copy()

    hoja2 = df_mat[['n_inscripto', 'nombre', 'apellido', 'dni', 'celular', 'certificado']].copy()
    hoja3 = df_cert[['n_inscripto', 'nombre', 'apellido', 'celular', 'apellido_nombre', 'presente']].copy()

    df_ins_con_mat = df_ins.merge(df_mat[['dni', 'certificado']], on='dni', how='left')
    resumen = df_ins_con_mat.groupby('partido').agg(
        inscriptos=('dni', 'count'),
        matriculados=('certificado', lambda x: x.notna().sum()),
        certificados=('certificado', lambda x: (x == 'Sí').sum())
    ).reset_index()

    total = pd.DataFrame({
        'partido': ['TOTAL'],
        'inscriptos': [resumen['inscriptos'].sum()],
        'matriculados': [resumen['matriculados'].sum()],
        'certificados': [resumen['certificados'].sum()]
    })
    resumen = pd.concat([resumen, total], ignore_index=True)

    with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
        hoja1.to_excel(writer, sheet_name='Inscriptos', index=False)
        hoja2.to_excel(writer, sheet_name='Matriculados', index=False)
        hoja3.to_excel(writer, sheet_name='Certificados', index=False)
        resumen.to_excel(writer, sheet_name='Consolidado', index=False)

    # Generar Word con gráficos
    word_name = f"{fecha_hoy}_INFORME_{safe_name}.docx"
    word_path = os.path.join(GENERATED_FOLDER, word_name)
    generar_word(df_ins, df_mat, df_cert, resumen, nombre_cap, word_path)

    resultado_final = {
        'excel': excel_name,
        'word': word_name,
        'inscriptos': len(df_ins),
        'matriculados': len(df_mat),
        'certificados': (df_mat['certificado'] == 'Sí').sum(),
        'resumen': resumen.to_dict('records')
    }
    del df_ins, df_mat, df_cert, df_inscriptos_raw, df_mat_cert_raw, df_ins_con_mat
    gc.collect()
    return resultado_final


def generar_informe_periodo(caps, periodo_label, tipo, word_path, excel_path):
    """Genera Excel + Word consolidando varias capacitaciones de un mismo período."""
    # --- Tabla detalle por capacitación ---
    detalle = pd.DataFrame([{
        'Capacitación': c['nombre'],
        'Fecha': c.get('fecha_realizacion', c['fecha']),
        'Inscriptos': c['inscriptos'],
        'Matriculados': c['matriculados'],
        'Certificados': c['certificados']
    } for c in caps])

    total_insc = detalle['Inscriptos'].sum()
    total_mat = detalle['Matriculados'].sum()
    total_cert = detalle['Certificados'].sum()

    # --- Consolidado geográfico (suma de los 'resumen' guardados por capacitación) ---
    geo = {}
    for c in caps:
        for fila in c.get('resumen', []):
            partido = fila.get('partido')
            if partido == 'TOTAL' or not partido:
                continue
            geo.setdefault(partido, {'inscriptos': 0, 'matriculados': 0, 'certificados': 0})
            geo[partido]['inscriptos'] += fila.get('inscriptos', 0)
            geo[partido]['matriculados'] += fila.get('matriculados', 0)
            geo[partido]['certificados'] += fila.get('certificados', 0)
    df_geo = pd.DataFrame([{'Partido': k, **v} for k, v in geo.items()]).sort_values('inscriptos', ascending=False) if geo else pd.DataFrame(columns=['Partido', 'inscriptos', 'matriculados', 'certificados'])

    resumen_gral = pd.DataFrame([{
        'Capacitaciones realizadas': len(caps),
        'Total inscriptos': total_insc,
        'Total matriculados': total_mat,
        'Total certificados': total_cert,
        '% Matriculación': round(total_mat / total_insc * 100, 1) if total_insc else 0,
        '% Certificación (s/matriculados)': round(total_cert / total_mat * 100, 1) if total_mat else 0
    }])

    # --- Excel ---
    with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
        resumen_gral.to_excel(writer, sheet_name='Resumen General', index=False)
        detalle.to_excel(writer, sheet_name='Detalle Capacitaciones', index=False)
        df_geo.to_excel(writer, sheet_name='Geográfico Consolidado', index=False)

    # --- Word ---
    doc = Document()
    title = doc.add_heading(f'INFORME {tipo.upper()} DE CAPACITACIONES', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'Período: {periodo_label}')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 117, 182)

    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha_p.add_run(f'Generado el {datetime.now().strftime("%d/%m/%Y")}').font.size = Pt(10)
    doc.add_paragraph()

    doc.add_heading('1. RESUMEN EJECUTIVO', level=1)
    doc.add_paragraph(
        f"""Durante el período {periodo_label} se realizaron {len(caps)} capacitaciones:
• Inscriptos totales: {total_insc} participantes
• Matriculados: {total_mat} participantes ({resumen_gral['% Matriculación'][0]}% de inscriptos)
• Certificados: {total_cert} participantes ({resumen_gral['% Certificación (s/matriculados)'][0]}% de matriculados)"""
    )

    doc.add_heading('2. CAPACITACIONES DEL PERÍODO', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Capacitación', 'Inscriptos', 'Matriculados', 'Certificados'
    for c in caps:
        row = table.add_row().cells
        row[0].text = str(c['nombre'])
        row[1].text = str(c['inscriptos'])
        row[2].text = str(c['matriculados'])
        row[3].text = str(c['certificados'])
    doc.add_paragraph()

    # Gráfico evolución por capacitación
    doc.add_heading('3. EVOLUCIÓN DE INSCRIPTOS POR CAPACITACIÓN', level=1)
    fig, ax = plt.subplots(figsize=(7, 3.8))
    ax.bar(detalle['Capacitación'], detalle['Inscriptos'], color='#2E75B6')
    ax.set_ylabel('Inscriptos')
    ax.set_title(f'Inscriptos por capacitación — {periodo_label}')
    plt.xticks(rotation=30, ha='right')
    plt.tight_layout()
    graf1 = os.path.join(GENERATED_FOLDER, 'temp_graf_periodo1.png')
    fig.savefig(graf1, dpi=90, bbox_inches='tight')
    plt.close(fig)
    plt.close('all')
    gc.collect()
    doc.add_picture(graf1, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Gráfico geográfico consolidado
    if not df_geo.empty:
        doc.add_heading('4. DISTRIBUCIÓN GEOGRÁFICA CONSOLIDADA', level=1)
        fig2, ax2 = plt.subplots(figsize=(7, 4.2))
        ax2.barh(df_geo['Partido'], df_geo['inscriptos'], color='#70AD47')
        ax2.set_xlabel('Inscriptos')
        ax2.set_title(f'Distribución geográfica consolidada — {periodo_label}')
        plt.tight_layout()
        graf2 = os.path.join(GENERATED_FOLDER, 'temp_graf_periodo2.png')
        fig2.savefig(graf2, dpi=90, bbox_inches='tight')
        plt.close(fig2)
        plt.close('all')
        gc.collect()
        doc.add_picture(graf2, width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        os.remove(graf2)

    doc.save(word_path)
    if os.path.exists(graf1):
        os.remove(graf1)

def generar_word(df_ins, df_mat, df_cert, resumen, nombre_cap, word_path):
    """Genera el informe en Word con gráficos."""
    doc = Document()

    # Título
    title = doc.add_heading('INFORME DE CAPACITACIÓN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(nombre_cap)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 117, 182)

    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha_p.add_run(f'Fecha: {datetime.now().strftime("%d/%m/%Y")}').font.size = Pt(10)
    doc.add_paragraph()

    # Resumen ejecutivo
    doc.add_heading('1. RESUMEN EJECUTIVO', level=1)
    total_insc = len(df_ins)
    total_mat = len(df_mat)
    total_cert = (df_mat['certificado'] == 'Sí').sum()

    resumen_texto = f"""Datos clave de la capacitación:
• Inscriptos totales: {total_insc} participantes
• Matriculados: {total_mat} participantes ({total_mat/total_insc*100:.1f}% de inscriptos)
• Certificados: {total_cert} participantes ({total_cert/total_mat*100:.1f}% de matriculados)
• Tasa de certificación global: {total_cert/total_insc*100:.1f}% sobre inscriptos"""
    doc.add_paragraph(resumen_texto)

    # Gráfico y análisis geográfico
    doc.add_heading('2. DISTRIBUCIÓN GEOGRÁFICA', level=1)

    resumen_graf = resumen[resumen['partido'] != 'TOTAL'].copy()
    fig, ax = plt.subplots(figsize=(7, 4.2))
    colores = ['#2E75B6', '#70AD47', '#FFC000', '#ED7D31', '#5B9BD5']
    bars = ax.barh(resumen_graf['partido'], resumen_graf['inscriptos'], color=colores[:len(resumen_graf)])
    ax.set_xlabel('Cantidad de Inscriptos')
    ax.set_title('Distribución Geográfica de Inscriptos por Partido')
    for bar, val in zip(bars, resumen_graf['inscriptos']):
        ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2, str(int(val)), va='center')
    plt.tight_layout()

    graf_path = os.path.join(GENERATED_FOLDER, 'temp_graf_geo.png')
    fig.savefig(graf_path, dpi=90, bbox_inches='tight')
    plt.close(fig)
    plt.close('all')
    gc.collect()

    doc.add_picture(graf_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Análisis:', level=2)
    analisis = f"""La capacitación mostró una concentración geográfica predominante en San Miguel 
con {resumen_graf[resumen_graf['partido']=='San Miguel']['inscriptos'].values[0] if 'San Miguel' in resumen_graf['partido'].values else 0} inscriptos, representando la mayoría del total. 
Otros partidos del GBA aportaron participación minoritaria pero significativa."""
    doc.add_paragraph(analisis)

    # Guardar
    doc.save(word_path)
    if os.path.exists(graf_path):
        os.remove(graf_path)
    gc.collect()

# ============================================================
# RUTAS DE LA WEB
# ============================================================

def cargar_datos():
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {'capacitaciones': []}

def guardar_datos(data):
    with open(DATA_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    if 'logged_in' not in session:
        return redirect(url_for('login'))
    return redirect(url_for('dashboard'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = hashlib.sha256(request.form['password'].encode()).hexdigest()
        if username == USUARIO['username'] and password == USUARIO['password']:
            session['logged_in'] = True
            return redirect(url_for('dashboard'))
        flash('Usuario o contraseña incorrectos', 'error')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('logged_in', None)
    return redirect(url_for('login'))

@app.route('/dashboard')
def dashboard():
    if 'logged_in' not in session:
        return redirect(url_for('login'))
    data = cargar_datos()
    total_cap = len(data['capacitaciones'])
    total_insc = sum(c.get('inscriptos', 0) for c in data['capacitaciones'])
    total_mat = sum(c.get('matriculados', 0) for c in data['capacitaciones'])
    total_cert = sum(c.get('certificados', 0) for c in data['capacitaciones'])
    return render_template('dashboard.html', total_cap=total_cap, total_insc=total_insc, 
                          total_mat=total_mat, total_cert=total_cert, capacitaciones=data['capacitaciones'])

@app.route('/nueva', methods=['GET', 'POST'])
def nueva_capacitacion():
    if 'logged_in' not in session:
        return redirect(url_for('login'))

    if request.method == 'POST':
        nombre = request.form['nombre']
        fecha_realizacion = request.form.get('fecha_realizacion') or datetime.now().strftime('%Y-%m-%d')
        inscriptos_num = int(request.form['inscriptos'])
        matriculados_num = int(request.form['matriculados'])
        certificados_num = int(request.form['certificados'])

        file_insc = request.files['file_inscriptos']
        file_mat = request.files['file_matriculados']

        if file_insc and allowed_file(file_insc.filename) and file_mat and allowed_file(file_mat.filename):
            filename_insc = secure_filename(file_insc.filename)
            filename_mat = secure_filename(file_mat.filename)
            filepath_insc = os.path.join(UPLOAD_FOLDER, filename_insc)
            filepath_mat = os.path.join(UPLOAD_FOLDER, filename_mat)
            file_insc.save(filepath_insc)
            file_mat.save(filepath_mat)

            try:
                resultado = procesar_capacitacion(filepath_insc, filepath_mat, nombre, 
                                                   inscriptos_num, matriculados_num, certificados_num)

                # Guardar en historial
                data = cargar_datos()
                cap_data = {
                    'id': len(data['capacitaciones']) + 1,
                    'nombre': nombre,
                    'fecha': datetime.now().strftime('%d/%m/%Y %H:%M'),
                    'fecha_realizacion': fecha_realizacion,
                    'inscriptos': resultado['inscriptos'],
                    'matriculados': resultado['matriculados'],
                    'certificados': resultado['certificados'],
                    'resumen': resultado['resumen'],
                    'excel': resultado['excel'],
                    'word': resultado['word']
                }
                data['capacitaciones'].append(cap_data)
                guardar_datos(data)

                flash('Capacitación procesada exitosamente', 'success')
                return redirect(url_for('ver_capacitacion', id=cap_data['id']))

            except Exception as e:
                flash(f'Error al procesar: {str(e)}', 'error')
        else:
            flash('Archivos no válidos. Solo se permiten .xlsx y .xls', 'error')

    return render_template('nueva.html')

@app.route('/capacitacion/<int:id>')
def ver_capacitacion(id):
    if 'logged_in' not in session:
        return redirect(url_for('login'))
    data = cargar_datos()
    cap = next((c for c in data['capacitaciones'] if c['id'] == id), None)
    if not cap:
        flash('Capacitación no encontrada', 'error')
        return redirect(url_for('dashboard'))
    return render_template('detalle.html', cap=cap)

@app.route('/informes', methods=['GET', 'POST'])
def informes():
    if 'logged_in' not in session:
        return redirect(url_for('login'))

    data = cargar_datos()
    caps = data['capacitaciones']

    # Períodos disponibles según las fechas reales de las capacitaciones cargadas
    periodos_disponibles = {'trimestral': set(), 'semestral': set(), 'anual': set()}
    for c in caps:
        try:
            f = datetime.strptime(c.get('fecha_realizacion', c['fecha'][:10]), '%Y-%m-%d')
        except ValueError:
            f = datetime.strptime(c['fecha'].split(' ')[0], '%d/%m/%Y')
        for tipo in periodos_disponibles:
            periodos_disponibles[tipo].add(obtener_periodo(f, tipo))

    resultado_generado = None
    if request.method == 'POST':
        tipo = request.form['tipo']
        periodo_label = request.form['periodo']

        caps_filtradas = []
        for c in caps:
            try:
                f = datetime.strptime(c.get('fecha_realizacion', c['fecha'][:10]), '%Y-%m-%d')
            except ValueError:
                f = datetime.strptime(c['fecha'].split(' ')[0], '%d/%m/%Y')
            if obtener_periodo(f, tipo) == periodo_label:
                caps_filtradas.append(c)

        if not caps_filtradas:
            flash('No hay capacitaciones cargadas para ese período', 'error')
        else:
            fecha_hoy = datetime.now().strftime('%Y%m%d')
            safe_periodo = periodo_label.replace(' ', '_')
            excel_name = f"{fecha_hoy}_INFORME_{tipo.upper()}_{safe_periodo}.xlsx"
            word_name = f"{fecha_hoy}_INFORME_{tipo.upper()}_{safe_periodo}.docx"
            excel_path = os.path.join(GENERATED_FOLDER, excel_name)
            word_path = os.path.join(GENERATED_FOLDER, word_name)
            generar_informe_periodo(caps_filtradas, periodo_label, tipo, word_path, excel_path)
            resultado_generado = {'excel': excel_name, 'word': word_name, 'periodo': periodo_label, 'cantidad': len(caps_filtradas)}
            flash('Informe generado exitosamente', 'success')

    periodos_disponibles = {k: sorted(v, reverse=True) for k, v in periodos_disponibles.items()}
    return render_template('informes.html', periodos=periodos_disponibles, resultado=resultado_generado)


@app.route('/descargar/<tipo>/<filename>')
def descargar(tipo, filename):
    if 'logged_in' not in session:
        return redirect(url_for('login'))
    filepath = os.path.join(GENERATED_FOLDER, filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    flash('Archivo no encontrado', 'error')
    return redirect(url_for('dashboard'))

if __name__ == '__main__':
    print("=" * 60)
    print("SISTEMA DE GESTIÓN DE CAPACITACIONES")
    print("=" * 60)
    print("Abriendo en: http://127.0.0.1:5000")
    print("Usuario: admin")
    print("Contraseña: admin123")
    print("=" * 60)
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=False, host='0.0.0.0', port=port)

if __name__ == '__main__':
    print("=" * 60)
    print("SISTEMA DE GESTIÓN DE CAPACITACIONES")
    print("=" * 60)
    print("Abriendo en: http://127.0.0.1:5000")
    print("Usuario: admin")
    print("Contraseña: admin123")
    print("=" * 60)
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=False, host='0.0.0.0', port=port)
